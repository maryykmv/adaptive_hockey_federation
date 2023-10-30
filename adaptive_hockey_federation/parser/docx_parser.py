import os
import re
from datetime import date
from typing import Optional

import docx  # type: ignore

from adaptive_hockey_federation.core.user_card import (  # type: ignore
    HockeyData,
)

NAME = '[И|и][М|м][Я|я]'
SURNAME = '[Ф|ф][А|а][М|м][И|и][Л|л][И|и][Я|я]'
PATRONYMIC = '[О|о][Т|т]?[Ч|ч][Е|е][С|с][Т|т][В|в][О|о]'
DATE_OF_BIRTH = '[Д|д][А|а][Т|т][А|а] [Р|р][О|о].+'
TEAM = '[К|к][О|о][М|м][А|а][Н|н][Д|д][А|а]'
PLAYER_NUMBER = '[И|и][Г|г][Р|р][О|о][В|в][О|о][Й|й]'
POSITION = '[П|п][О|о][З|з][И|и][Ц|ц][И|и][Я|я]'
NUMERIC_STATUS = '[Ч|ч].+[С|с][Т|т].+'
PLAYER_CLASS = '[К|к][Л|л][А|а][С|с][С|с]'


def read_file_columns(file: docx) -> list[docx]:
    """Функция находит таблицы в файле и возвращает список объектов
    docx с данными каждого столбца.
    """
    return [
        column
        for table in file.tables
        for index, column in enumerate(table.columns)
    ]


def read_file_text(file: docx) -> list[str]:
    """Функция находит текстовые данные в файле и возвращает список объектов
    docx с найденными данными.
    """
    return [
        run.text
        for paragraph in file.paragraphs
        for run in paragraph.runs
    ]


def columns_parser(
        columns: list[docx],
        regular_expression: str,
) -> list[Optional[str]]:
    """Функция находит столбец по названию и списком выводит содержимое
    каждой ячейки этого столбца.
    """
    output = [
        text if text
        else None
        for column in columns
        if re.search(
            regular_expression,
            list(cell.text for cell in column.cells)[0]
        )
        for text in list(cell.text for cell in column.cells)[1:]
    ]
    if not output:
        count = 0
        for column in columns:
            for index, cell in enumerate(column.cells):
                if re.search(r'п/п', cell.text):
                    for cell in column.cells[index + 1:]:
                        if cell.text and len(cell.text) < 4:
                            count += 1
                        else:
                            break
                else:
                    if count > 0:
                        break
        for column in columns:
            for index, cell in enumerate(column.cells):
                if re.search(regular_expression, cell.text):
                    for cell in column.cells[index + 1:index + 1 + count]:
                        output.append(cell.text)

    return output


def find_names(columns: list[docx], regular_expression: str) -> list[str]:
    """Функция парсит в искомом столбце имена. Опирается на шаблон ФИО
    (имя идет после фамилии на втором месте).
    """
    names_list = columns_parser(columns, regular_expression)
    return [
        name.split()[1].rstrip()
        for name in names_list
        if name
    ]


def find_surnames(columns: list[docx], regular_expression: str) -> list[str]:
    """Функция парсит в искомом столбце фамилии. Опирается на шаблон ФИО
    (фамилия идет на первом месте).
    """
    surnames_list = columns_parser(columns, regular_expression)
    return [
        surname.split()[0].rstrip()
        for surname in surnames_list
        if surname
    ]


def find_patronymics(
        columns: list[docx],
        regular_expression: str,
) -> list[str]:
    """Функция парсит в искомом столбце отчества. Опирается на шаблон ФИО
    (отчество идет на последнем месте).
    """
    patronymics_list = columns_parser(columns, regular_expression)
    return [
        patronymic.replace('/', ' ').split()[2].rstrip().rstrip(',')
        if patronymic and len(patronymic.split()) > 2
        else 'Отчество отсутствует'
        for patronymic in patronymics_list
    ]


def find_dates_of_birth(
        columns: list[docx],
        regular_expression: str,
) -> list[date]:
    """Функция парсит в искомом столбце дату рождения
    и опирается на шаблон дд.мм.гггг.
    """
    dates_of_birth_list = columns_parser(columns, regular_expression)
    dates_of_birth_list_clear = []
    for date_of_birth in dates_of_birth_list:
        if date_of_birth:
            try:
                for day, month, year in [
                    re.sub(r'\D', ' ', date_of_birth).split()
                ]:
                    if len(year) == 2:
                        if int(year) > 23:
                            year = '19' + year
                        else:
                            year = '20' + year
                    dates_of_birth_list_clear.append(
                        date(int(year), int(month), int(day))
                    )
            except ValueError or IndexError:  # type: ignore
                dates_of_birth_list_clear.append(date(1900, 1, 1))
        else:
            dates_of_birth_list_clear.append(date(1900, 1, 1))

    return dates_of_birth_list_clear


def find_team(
        text: list[str],
        columns: list[docx],
        regular_expression: str,
) -> str:
    """Функция парсит название команды.
    """
    text_clear = ' '.join(text)
    text_clear = re.sub(
        r'\W+|_+|ХК|СХК|ДЮСХК|Хоккейный клуб|по незрячему хоккею'
        '|по специальному хоккею|Спец хоккей|по специальному|по следж-хоккею',
        ' ',
        text_clear
    ).split()  # type: ignore
    try:
        return [
            'Молния Прикамья'
            if text_clear[index + 2] == 'Прикамья'
            else 'Ак Барс'
            if text_clear[index + 1] == 'Ак'
            else 'Снежные Барсы'
            if text_clear[index + 1] == 'Снежные'
            else 'Хоккей Для Детей'
            if text_clear[index + 1] == 'Хоккей'
            else 'Дети-Икс'
            if text_clear[index + 1] == 'Дети'
            else 'СКА-Стрела'
            if text_clear[index + 1] == 'СКА'
            else 'Сборная Новосибирской области'
            if text_clear[index + 2] == 'Новосибирской'
            else 'Атал'
            if text_clear[index + 3] == 'Атал'
            else 'Крылья Мечты'
            if text_clear[index + 2] == 'мечты'
            else 'Огни Магнитки'
            if text_clear[index + 1] == 'Огни'
            else 'Энергия Жизни Краснодар'
            if text_clear[index + 3] == 'Краснодар'
            else 'Энергия Жизни Сочи'
            if text_clear[index + 4] == 'Сочи'
            else 'Динамо-Москва'
            if text_clear[index + 1] == 'Динамо'
            else 'Крылья Советов'
            if text_clear[index + 2] == 'Советов'
            else 'Красная Ракета'
            if text_clear[index + 2] == 'Ракета'
            else 'Красная Молния'
            if text_clear[index + 2] == 'молния'
            else 'Сахалинские Львята'
            if text_clear[index + 1] == 'Сахалинские'
            else 'Мамонтята Югры'
            if text_clear[index + 1] == 'Мамонтята'
            else 'Уральские Волки'
            if text_clear[index + 1] == 'Уральские'
            else 'Нет названия команды'
            if text_clear[index + 1] == 'Всего'
            else text_clear[index + 1].capitalize()
            for index, txt in enumerate(text_clear)
            if re.search(regular_expression, txt)
        ][0]
    except IndexError:
        for column in columns:
            for cell in column.cells:
                if re.search(regular_expression, cell.text):
                    txt = re.sub(r'\W', ' ', cell.text)
                    return txt.split()[1].capitalize()

        return 'Название команды не найдено'


def find_players_number(
        columns: list[docx],
        regular_expression: str,
) -> list[int]:
    """Функция парсит в искомом столбце номер игрока.
    """
    players_number_list = columns_parser(columns, regular_expression)
    players_number_list_clear = []
    for player_number in players_number_list:
        if player_number:
            try:
                players_number_list_clear.append(
                    int(re.sub(r'\D', '', player_number)[:2])
                )
            except ValueError:
                players_number_list_clear.append(0)
        else:
            players_number_list_clear.append(0)

    return players_number_list_clear


def find_positions(columns: list[docx], regular_expression: str) -> list[str]:
    """Функция парсит в искомом столбце позицию игрока на поле.
    """
    positions_list = columns_parser(columns, regular_expression)
    return [
        'нападающий'
        if re.search(
            r'^н|^Н|^H|^Нп|^нл|^нп|^цн|^лн|^Нап|^№|^А,|^К,',
            position.lstrip()
        )
        else 'защитник'
        if re.search(r'^з|^З|^Зщ|^Защ', position.lstrip())
        else 'вратарь'
        if re.search(r'^Вр|^В|^вр', position.lstrip())
        else 'Позиция записана неверно'
        if not re.sub(r'\n|\(.+|\d', '', position)
        else re.sub(
            r'\n|\(.+|\d|Капитан',
            '',
            position
        ).lower().rstrip().replace(',', '').lstrip()
        for position in positions_list
        if position
    ]


def parser(file: docx) -> list[HockeyData]:
    """Функция собирает все данные об игроке
    и передает их в dataclass.
    """
    columns_from_file = read_file_columns(file)
    text_from_file = read_file_text(file)
    names = find_names(columns_from_file, NAME)
    surnames = find_surnames(columns_from_file, SURNAME)
    patronymics = find_patronymics(columns_from_file, PATRONYMIC)
    dates_of_birth = find_dates_of_birth(
        columns_from_file,
        DATE_OF_BIRTH,
    )
    team = find_team(text_from_file, columns_from_file, TEAM)
    players_number = find_players_number(columns_from_file, PLAYER_NUMBER)
    positions = find_positions(columns_from_file, POSITION)

    return [
        HockeyData(
            names[index],
            surnames[index],
            patronymics[index],
            dates_of_birth[index],
            team,
            players_number[index],
            positions[index],
        )
        for index in range(len(names))
    ]


if __name__ == '__main__':
    files_dir = '/Users/frost/dev/adaptive_hockey_federation/Именная заявка/'
    for root, directories, files in os.walk(files_dir):
        for file in files:
            if file.startswith('~'):
                pass
            else:
                if (
                    file.endswith('.docx')
                    and file != 'На мандатную комиссию.docx'
                    and file != (
                        'Именная заявка следж-хоккей Энергия Жизни Сочи.docx'
                    )
                    and file != 'ФАХ Сияжар Взрослые.docx'
                ):
                    parser(docx.Document(os.path.join(root, file)))
